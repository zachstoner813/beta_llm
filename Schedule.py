import os
import os.path
import pickle
from googleapiclient.discovery import build
from google_auth_oauthlib.flow import InstalledAppFlow
from google.auth.transport.requests import Request
from google.oauth2.credentials import Credentials
from docx import Document


# If modifying these SCOPES, delete the file token.pickle.
SCOPES = ['https://www.googleapis.com/auth/spreadsheets.readonly']

def main():
    creds = None
    # The file token.pickle stores the user's access and refresh tokens, and is
    # created automatically when the authorization flow completes for the first time.
    if os.path.exists('token.pickle'):
        with open('token.pickle', 'rb') as token:
            creds = pickle.load(token)
    # If there are no (valid) credentials available, let the user log in.
    if not creds or not creds.valid:
        if creds and creds.expired and creds.refresh_token:
            creds.refresh(Request())
        else:
            flow = InstalledAppFlow.from_client_secrets_file(
                '/Users/mac/Documents/Python Projects/API Projects/Beta Bot Segmented/client_secret.json', SCOPES)
            creds = flow.run_local_server(port=8081)
        # Save the credentials for the next run
        with open('token.pickle', 'wb') as token:
            pickle.dump(creds, token)

    service = build('sheets', 'v4', credentials=creds)

    # Call the Sheets API
    sheet = service.spreadsheets()
    result = sheet.values().get(spreadsheetId='1VMpk304hdNdmV6PfSKpPSLvVc543d54m_2rvhS-v53M',
                                range='Sheet1!A1:E').execute()
    values = result.get('values', [])

    if not values:
        print('No data found.')
    else:
        doc = Document()
        for row in values:
            row += [''] * (2 - len(row))  # Ensure row has 2 elements
            # Add a row in the Word document
            paragraph = doc.add_paragraph()
            paragraph.add_run('%s, %s, %s, %s' % (row[0], row[1], row[2], row[3]))
        
        # Ensure the directory exists
        directory = "var_files"
        if not os.path.exists(directory):
            os.makedirs(directory)

         # Define the file path
        file_path = os.path.join(directory, 'BAP_sheet_data.docx')
        
        # Check if the file already exists and delete it if it does
        if os.path.exists(file_path):
            os.remove(file_path)
        
        # Save the new document
        doc.save(file_path)


if __name__ == '__main__':
    main()
    

# packages required to run "pip install google-api-python-client google-auth-oauthlib google-auth python-docx"


